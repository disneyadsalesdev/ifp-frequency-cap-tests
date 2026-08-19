"""Update resume docx with nCPM bullet and export PDF."""
from __future__ import annotations

import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SOURCE = Path(
    r"c:\Users\syeda012\OneDrive - The Walt Disney Company\Documents\Resume\Ahmed_Syed_Resume(1).docx"
)
OUTPUT_DOCX = Path(
    r"c:\Users\syeda012\OneDrive - The Walt Disney Company\Documents\Resume\Ahmed_Syed_Resume(2).docx"
)
OUTPUT_PDF = Path(
    r"c:\Users\syeda012\OneDrive - The Walt Disney Company\Documents\Resume\Ahmed_Syed_Resume(2).pdf"
)

NEW_BULLET = (
    "Built an nCPM yield model normalizing CPM by ad duration to compare monetization "
    "efficiency across formats and content; used outputs to inform floor pricing and "
    "inventory allocation strategy."
)

# Insert after the CPM dashboard bullet (para index 20); before programmatic pricing bullet.
INSERT_AFTER_INDEX = 20


def insert_bullet_after(reference: Paragraph, text: str) -> Paragraph:
    """Clone bullet formatting from reference paragraph and insert after it."""
    new_p = deepcopy(reference._p)

    # Keep paragraph properties (style, numbering); remove content runs.
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)

    run_elem = OxmlElement("w:r")
    text_elem = OxmlElement("w:t")
    text_elem.set(qn("xml:space"), "preserve")
    text_elem.text = text
    run_elem.append(text_elem)
    new_p.append(run_elem)

    reference._p.addnext(new_p)
    return Paragraph(new_p, reference._parent)


def add_bullet_to_resume() -> tuple[int, str, str]:
    """Add bullet and save updated docx. Returns (insert_index, before_text, after_text)."""
    shutil.copy2(SOURCE, OUTPUT_DOCX)
    doc = Document(str(OUTPUT_DOCX))

    ref_para = doc.paragraphs[INSERT_AFTER_INDEX]
    before_text = doc.paragraphs[INSERT_AFTER_INDEX + 1].text.strip()

    new_para = insert_bullet_after(ref_para, NEW_BULLET)
    doc.save(str(OUTPUT_DOCX))

    return INSERT_AFTER_INDEX + 1, ref_para.text.strip(), before_text


def convert_to_pdf() -> str:
    """Try docx2pdf, LibreOffice, then win32com. Returns method used."""
    docx_path = str(OUTPUT_DOCX)
    pdf_path = str(OUTPUT_PDF)

    # 1. docx2pdf
    try:
        from docx2pdf import convert

        convert(docx_path, pdf_path)
        if OUTPUT_PDF.exists():
            return "docx2pdf"
    except Exception as exc:
        print(f"docx2pdf failed: {exc}", file=sys.stderr)

    # 2. LibreOffice headless
    import subprocess

    for soffice in ("soffice", "soffice.exe"):
        try:
            result = subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(OUTPUT_PDF.parent),
                    docx_path,
                ],
                capture_output=True,
                text=True,
                timeout=120,
            )
            if result.returncode == 0:
                # LibreOffice names output from input filename
                lo_pdf = OUTPUT_PDF.parent / f"{OUTPUT_DOCX.stem}.pdf"
                if lo_pdf.exists() and lo_pdf != OUTPUT_PDF:
                    lo_pdf.replace(OUTPUT_PDF)
                if OUTPUT_PDF.exists():
                    return "libreoffice"
        except (FileNotFoundError, subprocess.TimeoutExpired) as exc:
            print(f"LibreOffice ({soffice}) failed: {exc}", file=sys.stderr)

    # 3. win32com Word automation
    try:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # wdFormatPDF
        doc.Close()
        word.Quit()
        if OUTPUT_PDF.exists():
            return "win32com"
    except Exception as exc:
        print(f"win32com failed: {exc}", file=sys.stderr)

    raise RuntimeError("All PDF conversion methods failed")


def verify_docx() -> list[str]:
    """Re-extract Disney role bullets for verification."""
    doc = Document(str(OUTPUT_DOCX))
    bullets: list[str] = []
    in_disney_role = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if "Revenue & Yield Management Analyst" in text:
            in_disney_role = True
            continue
        if in_disney_role:
            if "Walt Disney Parks and Resorts" in text or "Marvel Entertainment" in text:
                break
            if p.style.name == "Heading 2" and p._p.pPr is not None:
                num_pr = p._p.pPr.find(qn("w:numPr"))
                if num_pr is not None:
                    bullets.append(text)
    return bullets


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(f"Source not found: {SOURCE}")

    insert_idx, prev_bullet, next_bullet = add_bullet_to_resume()
    pdf_method = convert_to_pdf()
    bullets = verify_docx()

    bullet_found = any(NEW_BULLET[:40] in b for b in bullets)
    bullet_position = None
    for i, b in enumerate(bullets):
        if NEW_BULLET[:40] in b:
            bullet_position = i + 1
            break

    print("STATUS: SUCCESS")
    print(f"DOCX: {OUTPUT_DOCX} (exists={OUTPUT_DOCX.exists()})")
    print(f"PDF:  {OUTPUT_PDF} (exists={OUTPUT_PDF.exists()})")
    print(f"PDF method: {pdf_method}")
    print(f"Inserted at paragraph index: {insert_idx}")
    print(f"Position among Disney role bullets: #{bullet_position} of {len(bullets)}")
    print(f"After: {prev_bullet[:80]}...")
    print(f"Before: {next_bullet[:80]}...")
    print(f"Bullet verified in docx: {bullet_found}")
    print("\n--- Disney role bullets ---")
    for i, b in enumerate(bullets, 1):
        marker = " <<< NEW" if NEW_BULLET[:40] in b else ""
        print(f"{i}. {b[:100]}{'...' if len(b) > 100 else ''}{marker}")


if __name__ == "__main__":
    main()
